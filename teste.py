from docx import Document

# Caminho do seu arquivo SQL
sql_file_path = r'C:\Users\bruno\Desktop\ALURA-SCRIPT.sql'

# Lê o conteúdo do arquivo SQL
with open(sql_file_path, 'r') as sql_file:
    sql_script = sql_file.read()

# Cria um novo documento do Word
doc = Document()

# Adiciona um título ao documento (opcional)
doc.add_heading('Script SQL', 0)

# Adiciona o conteúdo do script SQL
doc.add_paragraph(sql_script)

# Salva o documento em um arquivo .docx
doc.save(r'C:\Users\bruno\Desktop\script_sql.docx')

print('O script foi salvo em script_sql.docx')